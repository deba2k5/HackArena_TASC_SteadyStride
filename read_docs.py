import os
import docx
import pypdf
import pandas as pd

def extract_docx(file_path):
    print(f"=== Reading DOCX: {file_path} ===")
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            full_text.append(" | ".join(row_text))
    return "\n".join(full_text)

def extract_pdf(file_path):
    print(f"=== Reading PDF: {file_path} ===")
    reader = pypdf.PdfReader(file_path)
    full_text = []
    for page_num, page in enumerate(reader.pages):
        text = page.extract_text()
        full_text.append(f"--- Page {page_num+1} ---")
        full_text.append(text)
    return "\n".join(full_text)

def inspect_xlsx(file_path):
    print(f"=== Reading XLSX: {file_path} ===")
    xl = pd.ExcelFile(file_path)
    print("Sheets:", xl.sheet_names)
    for sheet_name in xl.sheet_names:
        df = xl.parse(sheet_name)
        print(f"\nSheet '{sheet_name}' shape: {df.shape}")
        print("Columns:", df.columns.tolist())
        print("First 3 rows:")
        print(df.head(3).to_string())

def main():
    base_dir = r"C:\Users\Debangshu05\Downloads\Supporting Documents - Touchless Agent\For Contestants"
    
    # 1. Read DOCX
    docx_path = os.path.join(base_dir, "AI_Invoicing_Challenge_Test_Data_Brief.docx")
    if os.path.exists(docx_path):
        text = extract_docx(docx_path)
        with open("AI_Invoicing_Challenge_Test_Data_Brief.txt", "w", encoding="utf-8") as f:
            f.write(text)
        print("Saved docx text to AI_Invoicing_Challenge_Test_Data_Brief.txt")
    
    # 2. Read PDF 1
    pdf1_path = os.path.join(base_dir, "TIA - Problem Statement.pdf")
    if os.path.exists(pdf1_path):
        text = extract_pdf(pdf1_path)
        with open("TIA_Problem_Statement.txt", "w", encoding="utf-8") as f:
            f.write(text)
        print("Saved pdf1 text to TIA_Problem_Statement.txt")
        
    # 3. Read PDF 2
    pdf2_path = os.path.join(base_dir, "TIA Hackathon Brief.pdf")
    if os.path.exists(pdf2_path):
        text = extract_pdf(pdf2_path)
        with open("TIA_Hackathon_Brief.txt", "w", encoding="utf-8") as f:
            f.write(text)
        print("Saved pdf2 text to TIA_Hackathon_Brief.txt")

    # 4. Inspect Excel
    xlsx_path = os.path.join(base_dir, "TASC_Sample_Database_vF.xlsx")
    if os.path.exists(xlsx_path):
        inspect_xlsx(xlsx_path)

if __name__ == "__main__":
    main()
